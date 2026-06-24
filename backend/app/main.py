# app/main.py

import io
import os
import csv
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from langchain_core.documents import Document

import docx
import pypdf

from app.config import ALLOW_OWN_KNOWLEDGE, PORT
from app.vector_store import get_collections, create_collection, get_vector_store, delete_collection
from app.faq_generator import (
    extract_faqs_from_text,
    extract_faqs_from_text_stream,
    expand_faq_questions,
    expand_faq_questions_stream,
    clean_rule_faqs,
    clean_rule_faqs_stream,
)
from app.retrieval import retrieve_and_rerank
from app.agents import rewrite_query, synthesize_answer
from app.rule_faq_generator import extract_faqs_rules
from app.language_utils import detect_language, get_db_language
from fastapi.responses import StreamingResponse
from qdrant_client import QdrantClient
from qdrant_client.models import Filter, FieldCondition, MatchValue
from app.config import QDRANT_URL

app = FastAPI(
    title="SAR Engine Manager & API Service",
    description="Admin UI backend and external integration gateway for the SAR FAQ system."
)

# Enable CORS for frontend API calls
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Schemas
class CreateCollectionRequest(BaseModel):
    name: str

class FAQItem(BaseModel):
    category: str
    question: str
    answer: str
    filename: Optional[str] = ""
    source_type: Optional[str] = "LLM"

class IngestRequest(BaseModel):
    collection_name: str
    filename: str
    faqs: List[FAQItem]
    language: str = "Thai"

class QueryRequest(BaseModel):
    query: str
    chat_history: List[Dict[str, str]] = []
    allow_own_knowledge: Optional[bool] = None

class AddFaqRequest(BaseModel):
    category: str = ""
    question: str
    answer: str
    language: str = "Thai"

class EditFaqRequest(BaseModel):
    # The original_question identifying the FAQ group to replace.
    original_question: str
    category: str = ""
    question: str
    answer: str
    language: str = "Thai"

class DeleteFaqRequest(BaseModel):
    original_question: str

@app.get("/api/v1/collections")
async def api_get_collections():
    """Lists Qdrant collections created within this app."""
    try:
        cols = get_collections()
        return {"collections": cols}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/stats")
async def api_get_collections_stats():
    """Returns basic stats for each valid collection in Qdrant."""
    try:
        valid_cols = get_collections()
        client = QdrantClient(url=QDRANT_URL)
        stats = []
        for col in valid_cols:
            try:
                info = client.get_collection(col)
                stats.append({
                    "name": col,
                    "points_count": info.points_count,
                    "status": info.status.value if hasattr(info.status, 'value') else str(info.status)
                })
            except Exception as e:
                print(f"[main] Error getting stats for collection {col}: {e}")
                continue
        return {"stats": stats}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/v1/collections/create")
async def api_create_collection(req: CreateCollectionRequest):
    """Creates a new Qdrant collection configured for hybrid search."""
    if not req.name.strip():
        raise HTTPException(status_code=400, detail="Collection name cannot be empty")
    try:
        create_collection(req.name)
        return {"status": "success", "collection": req.name.strip().lower()}
    except ValueError as ve:
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def _extract_text_from_bytes(file_bytes: bytes, ext: str) -> str:
    """Extracts clean plain text from an uploaded document's bytes."""
    if ext == "docx":
        doc = docx.Document(io.BytesIO(file_bytes))
        # Extract text paragraph by paragraph
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == "pdf":
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text_pages = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text()
            if t:
                text_pages.append(f"--- Page {i+1} ---\n{t}")
        return "\n".join(text_pages)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    return ""


@app.post("/api/v1/extract-faq")
async def api_extract_faq(
    file: UploadFile = File(...),
    language: str = Form("Thai"),
    num_questions: int = Form(10)
):
    """
    Extracts clean text from a document and generates FAQ pairs, streaming live
    progress back to the client as newline-delimited JSON (NDJSON).

    Each line is one JSON object. Event types:
        {"type": "progress", "stage": "extracting"|"cleaning", "current": i, "total": N}
        {"type": "status", "stage": "...", ...}
        {"type": "result", "filename": "...", "faqs": [...], "extraction_cost": float}
        {"type": "error", "detail": "..."}
    """
    filename = file.filename
    ext = filename.split(".")[-1].lower() if "." in filename else ""

    if ext not in ["docx", "pdf", "txt"]:
        raise HTTPException(status_code=400, detail=f"Unsupported file format: .{ext}. Supported formats are .docx, .pdf, .txt")

    # Read the upload before streaming starts (UploadFile is tied to the request).
    file_bytes = await file.read()

    def generate():
        def line(event: dict) -> str:
            return json.dumps(event, ensure_ascii=False) + "\n"

        try:
            # 1. Clean Text Extraction
            text = _extract_text_from_bytes(file_bytes, ext)
            if not text.strip():
                yield line({"type": "error", "detail": "No readable text found in document."})
                return
            yield line({"type": "status", "stage": "extracted_text", "chars": len(text)})

            # 2. FAQ Extraction using modular generator (streamed chunk by chunk)
            extracted_faqs_llm, extraction_cost = [], 0.0
            for ev in extract_faqs_from_text_stream(text, filename, language, num_questions):
                if ev["type"] == "extract_done":
                    extracted_faqs_llm = ev["faqs"]
                    extraction_cost = ev["cost"]
                else:
                    yield line(ev)

            # 3. Rule-based extraction + cleaning (streamed batch by batch)
            yield line({"type": "status", "stage": "rule_extraction"})
            extracted_faqs_rules = extract_faqs_rules(text, filename)

            cleaned_rules, cleaning_cost = [], 0.0
            for ev in clean_rule_faqs_stream(extracted_faqs_rules, language):
                if ev["type"] == "clean_done":
                    cleaned_rules = ev["faqs"]
                    cleaning_cost = ev["cost"]
                else:
                    yield line(ev)

            extracted_faqs = extracted_faqs_llm + cleaned_rules
            yield line({
                "type": "result",
                "filename": filename,
                "faqs": extracted_faqs,
                "extraction_cost": extraction_cost + cleaning_cost
            })
        except Exception as e:
            print(f"[main] Error during text extraction/FAQ generation: {e}")
            yield line({"type": "error", "detail": str(e)})

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # disable nginx proxy buffering so events flush live
        },
    )

@app.post("/api/v1/collections/ingest")
async def api_ingest_faqs(req: IngestRequest):
    """
    Embeds and saves the FAQ pairs to the chosen Qdrant collection, streaming live
    progress back to the client as newline-delimited JSON (NDJSON).

    Each line is one JSON object. Event types:
        {"type": "progress", "stage": "expanding"|"embedding", "current": i, "total": N}
        {"type": "status", "stage": "..."}
        {"type": "result", "status": "success", "count": int, "expansion_cost": float, "csv_path": "..."}
        {"type": "error", "detail": "..."}
    """
    if not req.faqs:
        raise HTTPException(status_code=400, detail="FAQ list is empty")

    # Number of documents pushed to Qdrant per add_documents call so embedding
    # progress can be streamed batch by batch instead of one opaque bulk call.
    EMBED_BATCH_SIZE = 50

    # Snapshot the request data before streaming starts.
    faqs_to_expand = []
    for faq in req.faqs:
        d = faq.model_dump() if hasattr(faq, 'model_dump') else faq.dict()
        d['original_question'] = faq.question
        faqs_to_expand.append(d)

    def generate():
        def line(event: dict) -> str:
            return json.dumps(event, ensure_ascii=False) + "\n"

        try:
            # 1. Expand FAQs (streamed question by question)
            expanded_faqs, expansion_cost = [], 0.0
            for ev in expand_faq_questions_stream(faqs_to_expand, language=req.language):
                if ev["type"] == "expand_done":
                    expanded_faqs = ev["faqs"]
                    expansion_cost = ev["cost"]
                else:
                    yield line(ev)

            # 2. Create documents
            docs = []
            for faq in expanded_faqs:
                page_content = f"Question: {faq['question']}\nAnswer: {faq['answer']}"
                metadata = {
                    "category": faq.get('category', ''),
                    "original_question": faq.get('original_question', faq.get('question', '')),
                    "answer": faq.get('answer', ''),
                    "source_file": faq.get('filename') or req.filename,
                    "source_type": faq.get('source_type', 'LLM')
                }
                docs.append(Document(page_content=page_content, metadata=metadata))

            # 3. Add to vector store in hybrid mode (streamed batch by batch)
            store = get_vector_store(req.collection_name)
            total = len(docs)
            for start in range(0, total, EMBED_BATCH_SIZE):
                batch = docs[start:start + EMBED_BATCH_SIZE]
                store.add_documents(batch)
                yield line({
                    "type": "progress",
                    "stage": "embedding",
                    "current": min(start + EMBED_BATCH_SIZE, total),
                    "total": total,
                })

            # 4. Export to CSV
            os.makedirs("exports", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            export_filename = f"exports/{req.collection_name}_{timestamp}.csv"

            with open(export_filename, mode='w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(["Category", "Filename", "Original Question", "Question Variation", "Answer", "Source Type"])
                for faq in expanded_faqs:
                    writer.writerow([
                        faq.get('category', ''),
                        faq.get('filename', req.filename),
                        faq.get('original_question', ''),
                        faq.get('question', ''),
                        faq.get('answer', ''),
                        faq.get('source_type', 'LLM')
                    ])

            yield line({
                "type": "result",
                "status": "success",
                "count": len(docs),
                "expansion_cost": expansion_cost,
                "csv_path": export_filename,
            })
        except Exception as e:
            print(f"[main] Error during ingestion: {e}")
            yield line({"type": "error", "detail": str(e)})

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # disable nginx proxy buffering so events flush live
        },
    )

@app.post("/api/v1/query/{collection_name}")
async def api_query_collection(collection_name: str, req: QueryRequest):
    """Exposes a dynamic SAR query endpoint for external application queries."""
    # Ensure collection exists
    collection_name = collection_name.strip().lower()
    valid_cols = get_collections()
    if collection_name not in valid_cols:
        raise HTTPException(status_code=404, detail=f"Collection '{collection_name}' not found or not initialized in this app")
        
    try:
        effective_allow_own_knowledge = (
            ALLOW_OWN_KNOWLEDGE
            if req.allow_own_knowledge is None
            else req.allow_own_knowledge
        )

        # 1. Detect Languages
        user_language = detect_language(req.query)
        db_language = get_db_language(collection_name)
        print(f"[Query] User Language: {user_language}, DB Language: {db_language}")

        # 2. Query Rewrite (translate to db_language)
        search_query = rewrite_query(req.query, req.chat_history, db_language)
        
        # 3. Retrieve & Rerank
        contexts, best_score = retrieve_and_rerank(collection_name, search_query)
        
        # 4. Write Answer (in user_language)
        response_text, match_type = synthesize_answer(
            req.query,
            contexts,
            best_score,
            allow_own_knowledge=effective_allow_own_knowledge,
            user_language=user_language
        )
        
        # Format sources
        sources = list(set([c["source_file"] for c in contexts]))
        
        # Format context for return JSON
        formatted_context = [
            {
                "content": c["content"],
                "source_file": c["source_file"],
                "category": c["category"]
            }
            for c in contexts
        ]
        
        return {
            "response": response_text,
            "match_type": match_type,
            "score": best_score,
            "allow_own_knowledge": effective_allow_own_knowledge,
            "sources": sources,
            "context": formatted_context
        }
    except Exception as e:
        print(f"[main] Error querying collection '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/export/{collection_name}")
async def api_export_collection(collection_name: str):
    """Exports all points in a collection to a CSV file."""
    try:
        client = QdrantClient(url=QDRANT_URL)
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Category", "Filename", "Original Question", "Question", "Answer", "Source Type"])
        
        offset = None
        while True:
            records, next_page = client.scroll(
                collection_name=collection_name,
                limit=100,
                offset=offset,
                with_payload=True
            )
            
            for record in records:
                payload = record.payload
                metadata = payload.get("metadata", {})
                
                cat = metadata.get("category", "")
                fname = metadata.get("source_file", "")
                orig_q = metadata.get("original_question", "")
                ans = metadata.get("answer", "")
                stype = metadata.get("source_type", "LLM")
                
                page_content = payload.get("page_content", "")
                lines = page_content.split("\n")
                q_var = ""
                for line in lines:
                    if line.startswith("Question: "):
                        q_var = line[len("Question: "):]
                        break
                        
                writer.writerow([cat, fname, orig_q, q_var, ans, stype])
                
            if next_page is None:
                break
            offset = next_page
            
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]), 
            media_type="text/csv", 
            headers={"Content-Disposition": f"attachment; filename={collection_name}_export.csv"}
        )
    except Exception as e:
        print(f"[main] Error exporting collection '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/v1/collections/{collection_name}/faqs")
async def api_get_collection_faqs(collection_name: str, limit: int = 100):
    """Fetches a sample of FAQs from a collection."""
    try:
        client = QdrantClient(url=QDRANT_URL)
        records, _ = client.scroll(
            collection_name=collection_name,
            limit=limit,
            with_payload=True
        )
        
        faqs = []
        for record in records:
            payload = record.payload
            metadata = payload.get("metadata", {})
            page_content = payload.get("page_content", "")
            
            q_var = ""
            lines = page_content.split("\n")
            for line in lines:
                if line.startswith("Question: "):
                    q_var = line[len("Question: "):]
                    break
                    
            faqs.append({
                "id": str(record.id),
                "category": metadata.get("category", ""),
                "original_question": metadata.get("original_question", ""),
                "question": q_var,
                "answer": metadata.get("answer", ""),
                "filename": metadata.get("source_file", ""),
                "source_type": metadata.get("source_type", "")
            })
            
        return {"faqs": faqs}
    except Exception as e:
        print(f"[main] Error getting FAQs for '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))


def _ensure_collection(collection_name: str) -> str:
    """Validates the collection is managed by this app; returns the normalized name."""
    collection_name = collection_name.strip().lower()
    if collection_name not in get_collections():
        raise HTTPException(status_code=404, detail=f"Collection '{collection_name}' not found or not initialized in this app")
    return collection_name


def _delete_faq_group(collection_name: str, original_question: str) -> None:
    """Deletes every point (all paraphrase variations) whose metadata.original_question matches."""
    client = QdrantClient(url=QDRANT_URL)
    client.delete(
        collection_name=collection_name,
        points_selector=Filter(
            must=[FieldCondition(key="metadata.original_question", match=MatchValue(value=original_question))]
        ),
    )


def _add_faq_group(collection_name: str, category: str, question: str, answer: str, language: str) -> tuple[int, float]:
    """Expands one FAQ into question variations, embeds them, and returns (count, expansion_cost)."""
    faq = {
        "category": category,
        "question": question,
        "answer": answer,
        "original_question": question,
        "source_type": "Manual",
    }
    expanded_faqs, expansion_cost = expand_faq_questions([faq], language=language)

    docs = []
    for f in expanded_faqs:
        page_content = f"Question: {f['question']}\nAnswer: {f['answer']}"
        metadata = {
            "category": f.get("category", ""),
            "original_question": f.get("original_question", question),
            "answer": f.get("answer", ""),
            "source_file": f.get("filename", "") or "manual",
            "source_type": f.get("source_type", "Manual"),
        }
        docs.append(Document(page_content=page_content, metadata=metadata))

    store = get_vector_store(collection_name)
    store.add_documents(docs)
    return len(docs), expansion_cost


@app.post("/api/v1/collections/{collection_name}/delete")
async def api_delete_collection(collection_name: str):
    """Deletes an entire Knowledge Base (Qdrant collection + registry entry)."""
    collection_name = _ensure_collection(collection_name)
    try:
        delete_collection(collection_name)
        return {"status": "success", "deleted": collection_name}
    except Exception as e:
        print(f"[main] Error deleting collection '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/collections/{collection_name}/faqs/add")
async def api_add_faq(collection_name: str, req: AddFaqRequest):
    """Adds a single new FAQ (expanded into variations and embedded) to a collection."""
    collection_name = _ensure_collection(collection_name)
    if not req.question.strip() or not req.answer.strip():
        raise HTTPException(status_code=400, detail="Question and answer are required")
    try:
        count, cost = _add_faq_group(collection_name, req.category, req.question, req.answer, req.language)
        return {"status": "success", "count": count, "expansion_cost": cost}
    except Exception as e:
        print(f"[main] Error adding FAQ to '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/collections/{collection_name}/faqs/edit")
async def api_edit_faq(collection_name: str, req: EditFaqRequest):
    """Replaces an existing FAQ group: deletes the old variations, then re-expands and embeds the new content."""
    collection_name = _ensure_collection(collection_name)
    if not req.question.strip() or not req.answer.strip():
        raise HTTPException(status_code=400, detail="Question and answer are required")
    try:
        _delete_faq_group(collection_name, req.original_question)
        count, cost = _add_faq_group(collection_name, req.category, req.question, req.answer, req.language)
        return {"status": "success", "count": count, "expansion_cost": cost}
    except Exception as e:
        print(f"[main] Error editing FAQ in '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/collections/{collection_name}/faqs/delete")
async def api_delete_faq(collection_name: str, req: DeleteFaqRequest):
    """Deletes an FAQ group (all paraphrase variations sharing the original question)."""
    collection_name = _ensure_collection(collection_name)
    try:
        _delete_faq_group(collection_name, req.original_question)
        return {"status": "success"}
    except Exception as e:
        print(f"[main] Error deleting FAQ from '{collection_name}': {e}")
        raise HTTPException(status_code=500, detail=str(e))
